"""
Word (.docx) generator for T-VER reports (PDD / Validation / Monitoring).

Structure follows the reference .docx templates (SomromScan branding, no
government seals — "จัดทำตามรูปแบบ T-VER"). All data comes from
report_data.report_data() — nothing here queries the DB directly.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_NAME = "TH Sarabun New"


def _set_run_font(run, size=11, bold=False):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)


def _para(doc, text="", size=11, bold=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        _set_run_font(run, size, bold)
    return p


def _heading(doc, text, size=13):
    return _para(doc, text, size=size, bold=True)


def _kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Cm(5)
    for label, value in rows:
        row = table.add_row()
        r1 = row.cells[0].paragraphs[0].add_run(label)
        _set_run_font(r1, 10, bold=True)
        r2 = row.cells[1].paragraphs[0].add_run(str(value))
        _set_run_font(r2, 10)
    doc.add_paragraph()
    return table


def _data_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        r = table.rows[0].cells[i].paragraphs[0].add_run(h)
        _set_run_font(r, 10, bold=True)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            r = row.cells[i].paragraphs[0].add_run(str(val))
            _set_run_font(r, 10)
    doc.add_paragraph()
    return table


def _signature_table(doc, labels):
    table = doc.add_table(rows=2, cols=len(labels))
    table.style = "Table Grid"
    for i, label in enumerate(labels):
        r = table.rows[0].cells[i].paragraphs[0].add_run(f"ลงชื่อ ..................... {label}")
        _set_run_font(r, 10)
        table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = table.rows[1].cells[i].paragraphs[0].add_run("( ..................... ) วันที่ ........./........./.........")
        _set_run_font(r2, 10)
        table.rows[1].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    return table


def _add_page_number_field(paragraph):
    run = paragraph.add_run()
    _set_run_font(run, 9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def _add_footer(doc, form_code):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{form_code} | SomromScan MRV Platform | หน้า ")
    _set_run_font(run, 9)
    _add_page_number_field(p)


def _header_block(doc, subtitle, title):
    _para(doc, "SomromScan — แพลตฟอร์ม MRV", size=13, bold=True)
    _para(doc, f"({subtitle})", size=9)
    doc.add_paragraph()
    _para(doc, title, size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, "โครงการลดก๊าซเรือนกระจกภาคสมัครใจตามมาตรฐานของประเทศไทย (T-VER)", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()


def _net_note(doc, ghg):
    if ghg.get("net_was_floored"):
        p = _para(doc, ghg["net_floor_note"], size=10)
        for run in p.runs:
            run.italic = True


def generate_pdd_docx(data: dict) -> Document:
    doc = Document()
    project = data["project"]
    _header_block(
        doc,
        "จัดทำตามรูปแบบเอกสารข้อเสนอโครงการ T-VER ขององค์การบริหารจัดการก๊าซเรือนกระจก",
        "เอกสารข้อเสนอโครงการ (Project Design Document: PDD)",
    )

    _kv_table(doc, [
        ("ชื่อโครงการ", project["name_th"] or project["name"]),
        ("ผู้พัฒนาโครงการ", project["developer_org"]),
        ("ประเภทโครงการ", project["forest_type_label"]),
        ("ฉบับที่ / วันที่", f"ฉบับที่ 1 / {data['generated_at_th']}"),
    ])

    _heading(doc, "๑. รายละเอียดทั่วไปของโครงการ")
    _kv_table(doc, [
        ("ที่ตั้งโครงการ", f"{project['district'] or ''} {project['province'] or ''}".strip()),
        ("พิกัดอ้างอิง", f"ละติจูด {project['latitude']}° เหนือ ลองจิจูด {project['longitude']}° ตะวันออก" if project["latitude"] else "ไม่มีข้อมูล"),
        ("ขนาดพื้นที่", f"{project['area_rai']:,.1f} ไร่ ({project['area_hectare']:,.2f} เฮกตาร์)" if project["area_rai"] else "-"),
        ("วันเริ่มดำเนินโครงการ", project["project_start_date_th"]),
        ("ระยะเวลาคิดเครดิต", f"{project['crediting_period_years']} ปี"),
    ])

    _heading(doc, "๒. คำอธิบายโครงการและกิจกรรมลดก๊าซเรือนกระจก")
    _para(doc, (
        f"โครงการมีวัตถุประสงค์เพื่ออนุรักษ์และฟื้นฟูพื้นที่{project['forest_type_label']} "
        "ด้วยการปลูกเสริมไม้ยืนต้น การดูแลรักษาต้นไม้ให้เจริญเติบโต และการป้องกันการตัดฟัน "
        "เพื่อเพิ่มการกักเก็บคาร์บอนในมวลชีวภาพเหนือและใต้พื้นดิน ควบคู่กับการสร้างรายได้จากคาร์บอนเครดิตให้แก่ชุมชน"
    ))

    _heading(doc, "๓. การประยุกต์ใช้ระเบียบวิธีและความเข้าเกณฑ์ (Eligibility & Additionality)")
    _para(doc, (
        f"โครงการประยุกต์ใช้ระเบียบวิธี {project['methodology']} ร่วมกับสมการแอลโลเมตริกของ Winrock "
        "ในการประเมินการกักเก็บคาร์บอน มีความเพิ่มเติม (additionality) เนื่องจากกิจกรรมการฟื้นฟูจะไม่เกิดขึ้น "
        "หากปราศจากแรงจูงใจจากคาร์บอนเครดิต"
    ))

    _heading(doc, "๔. ขอบเขตโครงการและแหล่งก๊าซเรือนกระจกที่เกี่ยวข้อง")
    _para(doc, (
        f"ขอบเขตโครงการครอบคลุมพื้นที่ {project['area_rai']:,.1f} ไร่ตามพิกัดที่ระบุ "
        "แหล่งกักเก็บที่พิจารณา ได้แก่ มวลชีวภาพเหนือพื้นดินและใต้พื้นดินของไม้ยืนต้น "
        "ก๊าซเรือนกระจกที่เกี่ยวข้องหลักคือคาร์บอนไดออกไซด์ (CO2)"
    ))

    _heading(doc, "๕. การกำหนดกรณีฐาน (Baseline Scenario)")
    _para(doc, (
        "กรณีฐานกำหนดจากสภาพการใช้ที่ดินเดิมก่อนมีโครงการ ซึ่งมีการกักเก็บคาร์บอนในระดับต่ำ "
        f"ปริมาณการกักเก็บกรณีฐานประเมินไว้เท่ากับ {data['ghg']['ghg_bsl']:,.1f} tCO2eq ต่อปี"
    ))

    _heading(doc, "๖. การประเมินปริมาณก๊าซเรือนกระจกที่คาดว่าจะลด/กักเก็บได้")
    _para(doc, f"ประมาณการปริมาณสุทธิตลอดระยะเวลาคิดเครดิต {project['crediting_period_years']} ปี")
    _data_table(
        doc,
        ["ปี", "ปริมาณสุทธิ (tCO2eq/ปี)", "สะสม (tCO2eq)"],
        [[r["year_be"], f"{r['net_tco2eq_year']:,.1f}", f"{r['cumulative_tco2eq']:,.1f}"] for r in data["yearly_projection"]],
    )

    _heading(doc, "๗. แผนการติดตามผล (Monitoring Plan)")
    _para(doc, (
        "โครงการจะติดตามพารามิเตอร์ ได้แก่ DBH จำนวนต้น ชนิดพันธุ์ และอัตราการรอด ปีละครั้ง "
        "ผ่านระบบ SomromScan โดยจัดทำรายงานการติดตามผล (Monitoring Report) ตามแบบ T-VER-S-F005-MR"
    ))

    _heading(doc, "๘. ผลกระทบด้านสิ่งแวดล้อมและการมีส่วนร่วมของผู้มีส่วนได้เสีย")
    _para(doc, (
        "โครงการก่อให้เกิดผลกระทบเชิงบวกต่อความหลากหลายทางชีวภาพและวิถีชีวิตชุมชน "
        "มีการรับฟังความคิดเห็นของผู้มีส่วนได้เสียในพื้นที่ก่อนดำเนินโครงการ และไม่พบผลกระทบเชิงลบที่มีนัยสำคัญ"
    ))

    _heading(doc, "๙. การรับรองโดยผู้พัฒนาโครงการ")
    _signature_table(doc, ["ผู้พัฒนาโครงการ", "ผู้มีอำนาจลงนาม"])

    _add_footer(doc, "T-VER-PDD")
    return doc


def generate_validation_docx(data: dict) -> Document:
    doc = Document()
    project = data["project"]
    _header_block(
        doc,
        "จัดทำโดยหน่วยตรวจสอบความใช้ได้และทวนสอบ VVB ตามรูปแบบ T-VER",
        "รายงานการตรวจสอบความใช้ได้ (Validation Report)",
    )

    _kv_table(doc, [
        ("ชื่อโครงการ", project["name_th"] or project["name"]),
        ("หน่วยตรวจสอบ (VVB)", data["vvb"]["name_th"]),
        ("เลขที่รายงาน / วันที่", f"VAL-{project['id']:04d} / {data['generated_at_th']}"),
        ("ระดับความเชื่อมั่น", "ระดับความเชื่อมั่นที่สมเหตุสมผล (Reasonable Assurance)"),
    ])

    _heading(doc, "๑. บทสรุปผลการตรวจสอบ (Validation Opinion)")
    _para(doc, (
        f"จากการตรวจสอบเอกสารข้อเสนอโครงการ (PDD) และการตรวจประเมิน หน่วยตรวจสอบมีความเห็นว่าโครงการ "
        f"{project['name_th'] or project['name']} จัดทำขึ้นสอดคล้องกับข้อกำหนดและระเบียบวิธีของกลไก T-VER "
        "มีการประเมินกรณีฐาน ความเพิ่มเติม ขอบเขต และแผนการติดตามที่เหมาะสม "
        "จึงเห็นควรให้ผ่านการตรวจสอบความใช้ได้เพื่อเสนอขอขึ้นทะเบียนต่อ อบก. ต่อไป"
    ))

    _heading(doc, "๒. ขอบเขต วัตถุประสงค์ และเกณฑ์การตรวจสอบ")
    _para(doc, (
        "การตรวจสอบมีวัตถุประสงค์เพื่อประเมินว่าเอกสารข้อเสนอโครงการเป็นไปตามระเบียบวิธีและข้อกำหนดของ T-VER หรือไม่ "
        "ครอบคลุมประเด็นด้านระเบียบวิธี กรณีฐาน ความเพิ่มเติม ขอบเขตโครงการ การประเมินปริมาณก๊าซเรือนกระจก "
        "และแผนการติดตาม โดยใช้เกณฑ์ตามเอกสารมาตรฐานของ อบก."
    ))

    _heading(doc, "๓. วิธีการตรวจสอบ")
    _para(doc, (
        "หน่วยตรวจสอบดำเนินการทบทวนเอกสาร (document review) ตรวจประเมินพื้นที่ (site visit) "
        "สัมภาษณ์ผู้พัฒนาโครงการและผู้มีส่วนได้เสีย และตรวจสอบความถูกต้องของข้อมูลและการคำนวณในระบบ SomromScan"
    ))

    _heading(doc, "๔. ผลการตรวจสอบตามประเด็นและข้อค้นพบ")
    _para(doc, "ประเภทข้อค้นพบ : CAR = ข้อบกพร่องที่ต้องแก้ไข, CL = ข้อขอความชัดเจน, FAR = ข้อพึงดำเนินการในอนาคต", size=9)
    if data["validation_findings"]:
        _data_table(
            doc,
            ["ประเภท", "ประเด็น/ข้อค้นพบ", "สถานะ"],
            [[f["type"], f["description"], f["status"]] for f in data["validation_findings"]],
        )
    else:
        p = _para(doc, data["validation_findings_note"], size=10)
        for run in p.runs:
            run.italic = True
        doc.add_paragraph()

    _heading(doc, "๕. ข้อสรุปการตรวจสอบความใช้ได้")
    _para(doc, (
        "ภายหลังการแก้ไขข้อบกพร่อง (CAR) และข้อขอความชัดเจน (CL) ครบถ้วนแล้ว หน่วยตรวจสอบสรุปว่าโครงการเป็นไปตาม "
        "ข้อกำหนดของกลไก T-VER และให้ความเห็นผ่านการตรวจสอบความใช้ได้ (Positive Validation Opinion)"
    ))

    _heading(doc, "๖. การรับรองโดยหน่วยตรวจสอบ (VVB)")
    _signature_table(doc, ["ผู้ตรวจสอบนำ (Lead Validator)", "ผู้อนุมัติทางเทคนิค"])

    _add_footer(doc, "T-VER-VR")
    return doc


def generate_monitoring_docx(data: dict) -> Document:
    doc = Document()
    project = data["project"]
    ghg = data["ghg"]
    _header_block(
        doc,
        "จัดทำตามรูปแบบรายงานการติดตามผล T-VER Monitoring Report ขององค์การบริหารจัดการก๊าซเรือนกระจก",
        "รายงานการติดตามผลการดำเนินโครงการ",
    )
    _para(doc, "Monitoring Report", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    _kv_table(doc, [
        ("ชื่อโครงการ", project["name_th"] or project["name"]),
        ("เลขที่ทะเบียนโครงการ", project["tgo_registration_number"]),
        ("รอบการติดตามที่", data["verification"]["cycle_number"]),
        ("วันที่จัดทำรายงาน", data["generated_at_th"]),
    ])

    _heading(doc, "ส่วนที่ ๑ ข้อมูลทั่วไปของโครงการ")
    _kv_table(doc, [
        ("ผู้พัฒนาโครงการ", project["developer_org"]),
        ("ประเภทโครงการ", project["forest_type_label"]),
        ("ที่ตั้งโครงการ", f"{project['district'] or ''} {project['province'] or ''}".strip()),
        ("พิกัดอ้างอิง", f"ละติจูด {project['latitude']}° เหนือ ลองจิจูด {project['longitude']}° ตะวันออก" if project["latitude"] else "ไม่มีข้อมูล"),
        ("ขนาดพื้นที่โครงการ", f"{project['area_rai']:,.1f} ไร่ ({project['area_hectare']:,.2f} เฮกตาร์)" if project["area_rai"] else "-"),
        ("ระเบียบวิธีที่ใช้", f"{project['methodology']} ร่วมกับสมการแอลโลเมตริกของ Winrock"),
        ("ระยะเวลาคิดเครดิต", f"{project['crediting_period_years']} ปี"),
    ])

    _heading(doc, "ส่วนที่ ๒ ขอบเขตการดำเนินงานและระเบียบวิธีการติดตาม")
    _para(doc, (
        f"การติดตามผลในรอบนี้ครอบคลุมพื้นที่โครงการทั้งหมด {project['area_rai']:,.1f} ไร่ "
        "โดยดำเนินการสำรวจและวัดขนาดต้นไม้ในแปลงตัวอย่างที่เป็นตัวแทนของพื้นที่ ตามหลักการวัดคาร์บอนภาคพื้นดินของ "
        "Winrock International คณะทำงานได้บันทึกค่าเส้นผ่านศูนย์กลางเพียงอก (DBH) ของต้นไม้แต่ละต้น "
        "ระบุชนิดพันธุ์ พิกัด และสถานะการรอดของต้นไม้ ผ่านระบบ SomromScan"
    ))

    _heading(doc, "ส่วนที่ ๓ พารามิเตอร์ที่ติดตามและวิธีการตรวจวัด")
    _data_table(
        doc,
        ["พารามิเตอร์", "หน่วย", "วิธีการตรวจวัด", "ความถี่", "แหล่งข้อมูล"],
        [
            ["เส้นผ่านศูนย์กลางเพียงอก (DBH)", "ซม.", "วัดด้วยเทปวัดเส้นรอบวงที่ระดับ 1.30 ม.", "ปีละครั้ง", "ภาคสนาม/IoT"],
            ["จำนวนต้นไม้", "ต้น", "นับจากการสำรวจและบันทึกในระบบ", "ปีละครั้ง", "SomromScan"],
            ["ชนิดพันธุ์ไม้", "-", "จำแนกชนิด/วงศ์โดยผู้เชี่ยวชาญ", "ปีละครั้ง", "ภาคสนาม"],
            ["อัตราการรอดของต้นไม้", "%", "สถานะมีชีวิต/ตาย ต่อจำนวนทั้งหมด", "ปีละครั้ง", "SomromScan"],
        ],
    )

    _heading(doc, "ส่วนที่ ๔ การคำนวณปริมาณก๊าซเรือนกระจก")
    _para(doc, "AGB = a × (DBH)^b [สมการแอลโลเมตริกของ Winrock]", size=10)
    _para(doc, "คาร์บอนรวม (tC) = (AGB + BGB) × 0.47", size=10)
    _para(doc, "ปริมาณคาร์บอนไดออกไซด์เทียบเท่า (tCO2eq) = คาร์บอนรวม × (44/12)", size=10)
    doc.add_paragraph()
    _data_table(
        doc,
        ["รายการ", "สัญลักษณ์", "ค่า (tCO2eq)"],
        [
            ["การกักเก็บของโครงการ (Project Removals)", "GHG_PROJ", f"{ghg['ghg_proj']:,.3f}"],
            ["การกักเก็บกรณีฐาน (Baseline Removals)", "GHG_BSL", f"{ghg['ghg_bsl']:,.3f}"],
            ["การรั่วไหล (Leakage)", "GHG_LK", f"{ghg['ghg_lk']:,.3f}"],
            ["ปริมาณสุทธิ (Net = PROJ − BSL − LK)", "GHG_NET", f"{ghg['ghg_net']:,.3f}"],
        ],
    )
    _net_note(doc, ghg)
    _para(doc, f"สรุป : ในรอบการติดตามนี้ โครงการมีปริมาณการกักเก็บก๊าซเรือนกระจกสุทธิเท่ากับ {ghg['ghg_net']:,.3f} ตันคาร์บอนไดออกไซด์เทียบเท่า (tCO2eq)")

    _heading(doc, "ส่วนที่ ๕ ผลการติดตามในรอบการติดตามนี้")
    _para(doc, f"ผลการสำรวจและคำนวณจำแนกตามช่วงชั้นขนาดเส้นผ่านศูนย์กลางเพียงอก (DBH) จำนวนต้นไม้ที่ทำการวัดในรอบนี้รวม {data['survival']['total_trees']:,} ต้น")
    _data_table(
        doc,
        ["ช่วงชั้น DBH", "จำนวนต้น", "กักเก็บ (tCO2eq)", "สัดส่วน (%)"],
        [[r["label_th"], r["tree_count"], f"{r['co2_tco2eq']:,.3f}", f"{r['pct']:.1f}"] for r in data["dbh_breakdown"]]
        + [["รวม", data["survival"]["total_trees"], f"{data['carbon']['total_co2_tonnes']:,.3f}", "100.0"]],
    )
    s = data["survival"]
    _para(doc, f"อัตราการรอดของต้นไม้ในรอบติดตาม : {s['survival_pct']}% (ต้นมีชีวิต {s['alive_count']} ต้น จากทั้งหมด {s['total_trees']} ต้น)")

    _heading(doc, "ส่วนที่ ๖ สรุปและการรับรอง")
    _para(doc, (
        "คณะผู้จัดทำขอรับรองว่าข้อมูลและผลการติดตามในรายงานฉบับนี้ถูกต้องตามความเป็นจริง จัดทำขึ้นตามระเบียบวิธีที่กำหนด "
        "และพร้อมรับการทวนสอบจากหน่วยตรวจสอบความใช้ได้และทวนสอบ (VVB) เพื่อเสนอต่อองค์การบริหารจัดการก๊าซเรือนกระจก (อบก.) ต่อไป"
    ))
    _signature_table(doc, ["ผู้พัฒนาโครงการ", "หน่วยตรวจสอบและทวนสอบ (VVB)", "เจ้าหน้าที่ อบก."])

    _heading(doc, "เอกสารอ้างอิง", size=11)
    _para(doc, "๑. องค์การบริหารจัดการก๊าซเรือนกระจก (องค์การมหาชน). ระเบียบวิธีการลดก๊าซเรือนกระจกภาคป่าไม้และการเกษตร (T-VER).", size=9)
    _para(doc, "๒. Winrock International. Standard Operating Procedures for Terrestrial Carbon Measurement.", size=9)

    _add_footer(doc, "T-VER-S-F005-MR")
    return doc


GENERATORS = {
    "pdd": generate_pdd_docx,
    "validation": generate_validation_docx,
    "monitoring": generate_monitoring_docx,
}


def generate_docx(data: dict) -> Document:
    report_type = data["report_type"]
    if report_type not in GENERATORS:
        raise ValueError(f"report_type ไม่ถูกต้อง: {report_type}")
    return GENERATORS[report_type](data)
